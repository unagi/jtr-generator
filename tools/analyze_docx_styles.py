"""DOCX style analyzer for ReportLab style design."""

from __future__ import annotations

import argparse
import json
import sys
from collections import Counter, defaultdict
from collections.abc import Iterable
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn


def _length_to_pt(value: Any) -> float | None:
    if value is None:
        return None
    if hasattr(value, "pt"):
        return float(value.pt)
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _normalize_hex(value: str | None) -> str | None:
    if not value:
        return None
    lowered = value.lower()
    if lowered == "auto":
        return "auto"
    if len(value) == 6:
        return f"#{value.upper()}"
    return f"#{value}"


def _rgb_to_hex(rgb: Any) -> str | None:
    if rgb is None:
        return None
    return _normalize_hex(str(rgb))


def _enum_name(value: Any) -> str | None:
    if value is None:
        return None
    return getattr(value, "name", str(value))


def _extract_shading(element: Any | None) -> dict[str, str] | None:
    if element is None:
        return None
    shading = element.xpath("./w:shd")
    if not shading:
        return None
    shd = shading[0]
    return {
        "fill": _normalize_hex(shd.get(qn("w:fill"))),
        "color": _normalize_hex(shd.get(qn("w:color"))),
        "value": shd.get(qn("w:val")),
    }


def _extract_paragraph_borders(element: Any | None) -> dict[str, Any] | None:
    if element is None:
        return None
    borders = element.xpath("./w:pBdr")
    if not borders:
        return None
    p_borders = borders[0]
    sides = {}
    for side in ("top", "bottom", "left", "right", "between"):
        border = p_borders.find(qn(f"w:{side}"))
        if border is None:
            continue
        size = border.get(qn("w:sz"))
        size_pt = float(size) / 8 if size and size.isdigit() else None
        sides[side] = {
            "value": border.get(qn("w:val")),
            "color": _normalize_hex(border.get(qn("w:color"))),
            "size_pt": size_pt,
            "space": border.get(qn("w:space")),
        }
    return sides or None


def _infer_role(alignment: str | None, max_font_size: float | None) -> str:
    if max_font_size is None:
        return "Body"
    if alignment == "CENTER" and max_font_size >= 18:
        return "Title"
    if max_font_size >= 13:
        return "Heading1"
    if max_font_size >= 12:
        return "Heading2"
    return "Body"


def _extract_paragraph(
    paragraph: Any,
    index: int,
    style_lookup: dict[str, str],
) -> dict[str, Any]:
    style = paragraph.style
    p_format = paragraph.paragraph_format
    p_pr = paragraph._p.pPr
    style_name = getattr(style, "name", None)
    style_id = getattr(style, "style_id", None)
    if p_pr is not None:
        p_style = p_pr.find(qn("w:pStyle"))
        if p_style is not None:
            inferred_id = p_style.get(qn("w:val"))
            if style_id is None:
                style_id = inferred_id
            if style_name is None and inferred_id in style_lookup:
                style_name = style_lookup[inferred_id]
    run_sizes = [r.font.size.pt for r in paragraph.runs if r.font.size]
    run_colors = [_rgb_to_hex(getattr(r.font.color, "rgb", None)) for r in paragraph.runs]
    run_colors = [color for color in run_colors if color]
    max_font_size = max(run_sizes) if run_sizes else None
    entry: dict[str, Any] = {
        "index": index,
        "style_name": style_name,
        "style_id": style_id,
        "text": paragraph.text.strip(),
        "alignment": _enum_name(p_format.alignment),
        "font_size_pt": max_font_size,
        "text_color": _most_common(run_colors),
        "spacing": {
            "before_pt": _length_to_pt(p_format.space_before),
            "after_pt": _length_to_pt(p_format.space_after),
            "line_spacing": _length_to_pt(p_format.line_spacing),
        },
        "indent": {
            "left_pt": _length_to_pt(p_format.left_indent),
            "right_pt": _length_to_pt(p_format.right_indent),
            "first_line_pt": _length_to_pt(p_format.first_line_indent),
        },
        "borders": _extract_paragraph_borders(p_pr),
        "shading": _extract_shading(p_pr),
        "runs": [],
    }
    entry["role"] = _infer_role(entry["alignment"], max_font_size)

    for run_index, run in enumerate(paragraph.runs):
        font = run.font
        run_entry = {
            "index": run_index,
            "text": run.text.strip(),
            "bold": font.bold,
            "italic": font.italic,
            "underline": font.underline,
            "font_name": font.name,
            "font_size_pt": _length_to_pt(font.size),
            "color": _rgb_to_hex(getattr(font.color, "rgb", None)),
            "highlight": _enum_name(getattr(font, "highlight_color", None)),
            "style_name": getattr(run.style, "name", None),
        }
        entry["runs"].append(run_entry)

    return entry


def _extract_table(table: Any, index: int) -> dict[str, Any]:
    style = table.style
    rows = len(table.rows)
    cols = len(table.columns)
    cells: list[list[dict[str, Any]]] = []
    for row in table.rows:
        row_cells: list[dict[str, Any]] = []
        for cell in row.cells:
            tc_pr = cell._tc.tcPr
            shading = _extract_shading(tc_pr) if tc_pr is not None else None
            row_cells.append({"text": cell.text.strip(), "shading": shading})
        cells.append(row_cells)

    return {
        "index": index,
        "style_name": getattr(style, "name", None),
        "rows": rows,
        "cols": cols,
        "cells": cells,
    }


def _collect_style_stats(paragraphs: list[dict[str, Any]]) -> dict[str, Any]:
    stats: dict[str, dict[str, list[Any]]] = defaultdict(lambda: defaultdict(list))
    examples: dict[str, list[str]] = defaultdict(list)

    for paragraph in paragraphs:
        style_name = paragraph.get("style_name") or paragraph.get("role") or "Unknown"
        stats[style_name]["alignment"].append(paragraph.get("alignment"))
        spacing = paragraph.get("spacing", {})
        indent = paragraph.get("indent", {})
        for key in ("before_pt", "after_pt", "line_spacing"):
            value = spacing.get(key)
            if value is not None:
                stats[style_name][key].append(value)
        for key in ("left_pt", "right_pt", "first_line_pt"):
            value = indent.get(key)
            if value is not None:
                stats[style_name][key].append(value)

        runs = paragraph.get("runs", [])
        for run in runs:
            for key in ("bold", "italic", "underline", "font_name", "font_size_pt", "color"):
                value = run.get(key)
                if value is not None and value != "":
                    stats[style_name][key].append(value)

        text = paragraph.get("text", "")
        if text and len(examples[style_name]) < 3:
            examples[style_name].append(text[:120])

    reportlab_styles: dict[str, Any] = {}
    for style_name, values in stats.items():
        reportlab_styles[style_name] = {
            "alignment": _most_common(values.get("alignment")),
            "font_name": _most_common(values.get("font_name")),
            "font_size_pt": _median(values.get("font_size_pt")),
            "text_color": _most_common(values.get("color")),
            "bold": _most_common(values.get("bold")),
            "italic": _most_common(values.get("italic")),
            "underline": _most_common(values.get("underline")),
            "space_before_pt": _median(values.get("before_pt")),
            "space_after_pt": _median(values.get("after_pt")),
            "line_spacing": _median(values.get("line_spacing")),
            "left_indent_pt": _median(values.get("left_pt")),
            "right_indent_pt": _median(values.get("right_pt")),
            "first_line_indent_pt": _median(values.get("first_line_pt")),
            "examples": examples.get(style_name, []),
        }

    return reportlab_styles


def _median(values: Iterable[float] | None) -> float | None:
    if not values:
        return None
    sorted_values = sorted(values)
    return sorted_values[len(sorted_values) // 2]


def _most_common(values: Iterable[Any] | None) -> Any | None:
    if not values:
        return None
    counter = Counter(values)
    return counter.most_common(1)[0][0]


def _extract_styles(document: Document) -> list[dict[str, Any]]:
    styles = []
    for style in document.styles:
        if style.type not in (WD_STYLE_TYPE.PARAGRAPH, WD_STYLE_TYPE.CHARACTER):
            continue
        font = style.font
        p_format = getattr(style, "paragraph_format", None)
        styles.append(
            {
                "name": style.name,
                "style_id": style.style_id,
                "type": style.type.name,
                "based_on": getattr(style.base_style, "name", None),
                "font": {
                    "name": font.name,
                    "size_pt": _length_to_pt(font.size),
                    "bold": font.bold,
                    "italic": font.italic,
                    "underline": font.underline,
                    "color": _rgb_to_hex(getattr(font.color, "rgb", None)),
                },
                "paragraph_format": None
                if p_format is None
                else {
                    "alignment": _enum_name(p_format.alignment),
                    "space_before_pt": _length_to_pt(p_format.space_before),
                    "space_after_pt": _length_to_pt(p_format.space_after),
                    "line_spacing": _length_to_pt(p_format.line_spacing),
                    "left_indent_pt": _length_to_pt(p_format.left_indent),
                    "right_indent_pt": _length_to_pt(p_format.right_indent),
                    "first_line_indent_pt": _length_to_pt(p_format.first_line_indent),
                },
            }
        )
    return styles


def analyze_docx(docx_path: Path) -> dict[str, Any]:
    document = Document(str(docx_path))
    styles = _extract_styles(document)
    style_lookup = {style["style_id"]: style["name"] for style in styles if style["style_id"]}
    paragraphs = [_extract_paragraph(p, i, style_lookup) for i, p in enumerate(document.paragraphs)]
    tables = [_extract_table(t, i) for i, t in enumerate(document.tables)]
    return {
        "source": str(docx_path),
        "paragraphs": paragraphs,
        "tables": tables,
        "styles": styles,
        "reportlab_styles": _collect_style_stats(paragraphs),
    }


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Analyze DOCX styles for ReportLab design.")
    parser.add_argument(
        "input",
        nargs="?",
        default="tests/fixtures/resume_sample.docx",
        help="DOCX file path to analyze",
    )
    parser.add_argument(
        "--output",
        default="skill/assets/data/career_sheet/docx_style_report.json",
        help="Output JSON path",
    )
    args = parser.parse_args(argv)

    input_path = Path(args.input)
    output_path = Path(args.output)
    if not input_path.exists():
        print(f"Input file not found: {input_path}", file=sys.stderr)
        return 1

    report = analyze_docx(input_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(json.dumps(report, ensure_ascii=True, indent=2), encoding="utf-8")
    print(f"DOCX analysis report written: {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
